#Se importan las librerias requeridas para la ejecución de la clase
import pandas as pd
import numpy as np
import os
import docx
from docx import Document
import shutil
from time import *

from datetime import datetime
import csv


#Se importa el .py de las interfaces gráficas
from Main_Window import *


#Se declara el .py de la clase de la ventana a utilizar
# y de la clase que procesa la información
from GraphicsMenu import *
from Graphic_Data import *

#Clase de la pantalla principal
class MainEpsolApp (QtWidgets.QMainWindow):
    def __init__(self, parent=None):
        super(MainEpsolApp,self).__init__(parent)
        self.ui = Ui_root()
        self.ui.setupUi(self)
        self.show()
        self.initial_dir = "../"
        self.makeDir()
        self.createdFiles = [] # Esta lista almacena todos los html generados para eliminarlos al cerrar la app 
        
        #Estas variables sirven para saber si hay ventanas abiertas y cerrarlas al cerrar la app
        self.availableMenu = False
        self.availableView = False
        self.availableCustomized = False

        
        #Variable que almacenará los archivos a utilizar en el programa
        self.files = None        
        # Variable que almacenará el merge de los archivos cargados
        self.data = None

        #Listeners de los botones de la interfaz principal
        self.ui.load_file_button.clicked.connect(self.loadFiles)
        self.ui.download_file_button.clicked.connect(self.download)
        self.ui.graphic_button.clicked.connect(self.availableGraphics)
        self.ui.report_button.clicked.connect(self.generateReport)
        self.ui.menuHelp.triggered.connect(self.info)
        

    #Declaracion de métodos relacionados a los botones de la interfaz principal

    #Método para cargar los datos a procesas   
    def loadFiles(self):
        self.files = QFileDialog.getOpenFileNames(self, "Abrir Archivos", self.initial_dir, "CSV Files (*.csv)") 
        if self.files[0]:
            graph = Graphic_Data()
            self.data = graph.merge(self.files[0])

    #Metodo que crea la carpeta nueva para almacenar
    # los archivos de la ejecucion
    def makeDir(self):
        
        nombre = datetime.now().strftime('ArchivosAnalizados_%H_%M_%d_%m_%Y') 
        self.path =self.initial_dir+nombre+"/"
        print(self.path)
        os.mkdir(self.path)

    #Devuelve el path de la carpeta generada para guardar los documentos
    #NOTA: EN ESTA VERSION AUNQUE SE LLAMA EN OTRAS CLASES NO SE OCUPA MAS QUE EN LAS
    #FUNCIONES QUE GENERAN LOS REPORTES
    def getPath(self):
        return self.path

    #Estos metodos reciben la señal si hay ventanas abiertas y su referencia
    def setAvailableMenu(self, available, screen):
        self.availableMenu = available
        self.son = screen
    def setAvailableView(self, available, screen):
        self.availableView = available
        self.grandson_1 = screen
    def setAvailableCustomized(self, available, screen):
        self.availableCustomized = available
        self.grandson_2 = screen

    #Este metodo recibe cada html generado y lo almacena en una lista
    def setFiles(self, files):
        doc = files
        self.createdFiles.append(doc)
        print("Lista de archivos: ")
        print(self.createdFiles)
        print("")
    
    #Metodo que descarga el csv de la informacion procesada
    def download(self):
    
        if self.data is None:
            aviso= "No existen archivos cargados"
            self.warning(anuncio=aviso)
        else:
            graph = Graphic_Data()
            self.data=graph.merge(self.files[0])
            self.download = pd.DataFrame(self.data)
            self.nombre = datetime.now().strftime('Dataframe_usado__%H_%M_%d_%m_%Y.csv')
            self.download.to_csv(str(self.nombre), header=True, index=False)
            self.archivo = shutil.copy(self.nombre,self.path)
            os.remove(self.nombre)
            nota = "Descarga de reporte CSV completa" 
            self.information(anuncio=nota)
            
    #Metodo para generar el reporte Word       
    def generateReport(self):
        if self.data is None:
            aviso= "No existen archivos cargados"
            self.warning(anuncio=aviso)
        else:
            path = self.path       
            lstFiles = []        
            lstDir = os.walk(path)        
    
            for root, dirs, files in lstDir:
                for fichero in files:
                    (nombreFichero, extension) = os.path.splitext(fichero)
                    if(extension == ".png"):
                        lstFiles.append(nombreFichero+extension)
            if len(lstFiles) == 0:
                aviso="No se ha generado ninguna gráfica"
                self.warning(anuncio=aviso)
            else:  
                for f in lstFiles:            
                    try:
                        doc = docx.Document(path+'\\'+'test.docx')
                        #doc.add_picture(str(path+'\\'+f),width=docx.shared.Inches(10), height=docx.shared.Cm(5))
                        doc.add_picture(str(path+'\\'+f))
                        doc.save(path+'\\'+'test.docx')
                    except Exception as e: 
            
                        print(e,"Generando archivo .docx")            
                        document = Document()
                        document.save(path+'\\'+'test.docx')
                        doc = docx.Document(path+'\\'+'test.docx')
                        doc.add_picture(str(path+'\\'+f))
                        
                        doc.save(path+'\\'+'test.docx') 
                nota = "Descarga de reporte WORD completa" 
                self.information(anuncio=nota)
                print("Archivo finalizado")        
                
            

    #Metodo que muestra el menu de las gráficas disponibles
    def availableGraphics(self):
        if self.files:
            if self.files[0]:
                #Se manda el merge que se genera con los datos cargados
                graph = Graphic_Data()
                self.data = graph.merge(self.files[0])
                #Se crea una instancia de GraphicsMenu
                self.menu = GraphicsMenu(parent=self)
                self.menu.cargarData(data=self.data)
                #Se establecen las opciones de gráficas que estan disponibles
                #A partir de los archivos cargados
                self.menu.comparar()
                #Se muestra interfaz
                self.menu.show()
               
            else:
                #Se muestra señal de emergencia
                aviso= "No existen archivos cargados"
                self.warning(anuncio=aviso)
        else:
            #Se muestra señal de emergencia
            aviso= "No existen archivos cargados"
            self.warning(anuncio=aviso)
    
    #Metodo que opera al cerrar la venta principal
    def closeEvent(self, event):
        if event:
            #Cierra las ventanas secundarias si estan abiertas
            for i in range(3):
                if self.availableMenu:
                    self.son.close()
                if self.availableView:
                    self.grandson_1.close()
                if self.availableCustomized:
                    self.grandson_2.close()
            #Borra los html generados durante la ejecucion
            for file in self.createdFiles:
                os.remove(file)
               
        
    #Metodo que miestra información de desarrollo
    def info(self):
        QMessageBox.about(self, "Informacion del programa",
                          "Programa Diseñado por:\nErick Rodriguez Orduña, \nCarlos " +
                          "Ramirez Rendon, \nNatalia Bonifaz "+
                          "Oviedo, \nMaría Juliana " +
                          "Pérez Barón, \nKatia Patricia " +
                          "Limon Fraga y \nJosé Emmanuel " + 
                          "De Los Santos Castro \nPara: EPSOL SA de CV")

    #Metodo que muestra el mensaje de advertencia sobre datos inexistentes
    def warning(self,anuncio):
        QMessageBox.warning(self, "Advertencia!",str(anuncio), buttons=QMessageBox.Ok)

    #Metodo que muestra el mensaje de aviso si se descarga el csv o se genera el word
    def information(self, anuncio):
        aviso = anuncio
        QMessageBox.information(self, "Aviso!",
                                        str(aviso))
#Main
if __name__ == "__main__":
    import sys, time, os
    from PyQt5.QtWidgets import QFileDialog, QMessageBox, QWidget
    from PyQt5.QtWebEngineWidgets import *
    from PyQt5.QtWidgets import *
    from PyQt5 import QtWidgets
    from PyQt5.QtCore import *

    app = QtWidgets.QApplication(sys.argv)
    my_app = MainEpsolApp()
    sys.exit(app.exec_())